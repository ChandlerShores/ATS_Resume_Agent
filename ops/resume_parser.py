"""Resume parsing utilities for extracting bullets from DOCX, PDF, and TXT files."""

import re
from pathlib import Path
from typing import Dict, Any, List, Optional

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import chardet
except ImportError:
    chardet = None

from ops.parsing_errors import (
    UnsupportedFormatError,
    FileReadError,
    NoBulletsFoundError
)
from ops.logging import logger


# Bullet point patterns
BULLET_PATTERNS = [
    r'^\s*[•●○◦▪▫■□‣⁃∙⦿⦾]',  # Unicode bullet chars
    r'^\s*[-–—]',  # Dashes
    r'^\s*\*',  # Asterisks
    r'^\s*\d+[\.)]\s+',  # Numbered lists: 1. or 1)
]

# Section headers that typically contain work experience bullets
EXPERIENCE_SECTION_PATTERNS = [
    r'experience',
    r'work\s+history',
    r'professional\s+experience',
    r'employment',
    r'career\s+history',
    r'relevant\s+experience',
]

# Patterns to filter out (not actual work bullets)
EXCLUDE_PATTERNS = [
    r'^(bachelor|master|phd|associate|b\.s\.|m\.s\.|ph\.d\.)',  # Education
    r'^(email|phone|address|linkedin|github):',  # Contact info
    r'^\d{3}[-.\s]?\d{3}[-.\s]?\d{4}',  # Phone numbers
    r'^[\w\.-]+@[\w\.-]+\.\w+',  # Email addresses
]


def parse_resume(file_path: str) -> Dict[str, Any]:
    """
    Parse resume file and extract bullets and metadata.
    
    Args:
        file_path: Path to resume file (.docx, .pdf, or .txt)
        
    Returns:
        Dict with keys:
            - bullets: List[str] - Extracted resume bullets
            - raw_text: str - Full text content
            - metadata: dict - File metadata (name, format, etc.)
            
    Raises:
        UnsupportedFormatError: If file format is not supported
        FileReadError: If file cannot be read
        NoBulletsFoundError: If no bullets could be extracted
    """
    file_path = Path(file_path)
    
    # Check file exists
    if not file_path.exists():
        raise FileReadError(f"File not found: {file_path}")
    
    # Check format
    ext = file_path.suffix.lower()
    if ext not in ['.docx', '.pdf', '.txt']:
        raise UnsupportedFormatError(
            f"Format {ext} not supported. Supported formats: .docx, .pdf, .txt"
        )
    
    # Extract text based on format
    try:
        if ext == '.docx':
            raw_text = extract_from_docx(file_path)
        elif ext == '.pdf':
            raw_text = extract_from_pdf(file_path)
        else:  # .txt
            raw_text = extract_from_txt(file_path)
    except Exception as e:
        raise FileReadError(f"Failed to read {file_path}: {str(e)}")
    
    if not raw_text or not raw_text.strip():
        raise FileReadError(f"No text content found in {file_path}")
    
    # Extract bullets
    bullets = extract_bullets(raw_text)
    
    if not bullets:
        raise NoBulletsFoundError(
            f"No bullets found in {file_path}. "
            "The resume may have non-standard formatting or bullets may be in images/tables."
        )
    
    # Warn if suspiciously few bullets
    if len(bullets) < 3:
        logger.warn(
            stage="parse_resume",
            msg=f"Only {len(bullets)} bullets found in {file_path.name}",
            file=str(file_path)
        )
    
    metadata = {
        "filename": file_path.name,
        "format": ext,
        "bullet_count": len(bullets),
        "char_count": len(raw_text),
    }
    
    return {
        "bullets": bullets,
        "raw_text": raw_text,
        "metadata": metadata
    }


def extract_from_docx(file_path: Path) -> str:
    """
    Extract text from DOCX file.
    
    Args:
        file_path: Path to DOCX file
        
    Returns:
        str: Extracted text
        
    Raises:
        FileReadError: If file cannot be read
    """
    if Document is None:
        raise FileReadError("python-docx not installed. Install with: pip install python-docx")
    
    try:
        doc = Document(file_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        
        # Also try to extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        
        return '\n'.join(paragraphs)
    
    except Exception as e:
        raise FileReadError(f"Failed to parse DOCX: {str(e)}")


def extract_from_pdf(file_path: Path) -> str:
    """
    Extract text from PDF file.
    
    Args:
        file_path: Path to PDF file
        
    Returns:
        str: Extracted text
        
    Raises:
        FileReadError: If file cannot be read
    """
    if pdfplumber is None:
        raise FileReadError("pdfplumber not installed. Install with: pip install pdfplumber")
    
    try:
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        
        return '\n'.join(text_parts)
    
    except Exception as e:
        raise FileReadError(f"Failed to parse PDF: {str(e)}")


def extract_from_txt(file_path: Path) -> str:
    """
    Extract text from TXT file with encoding detection.
    
    Args:
        file_path: Path to TXT file
        
    Returns:
        str: Extracted text
        
    Raises:
        FileReadError: If file cannot be read
    """
    # Try UTF-8 first
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        pass
    
    # Fallback to encoding detection
    if chardet is not None:
        try:
            with open(file_path, 'rb') as f:
                raw_data = f.read()
            detected = chardet.detect(raw_data)
            encoding = detected.get('encoding', 'utf-8')
            
            return raw_data.decode(encoding)
        except Exception as e:
            raise FileReadError(f"Failed to decode TXT file: {str(e)}")
    else:
        # Try common encodings
        for encoding in ['latin-1', 'windows-1252', 'iso-8859-1']:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        
        raise FileReadError(
            "Failed to detect file encoding. Install chardet for better detection: pip install chardet"
        )


def extract_bullets(text: str) -> List[str]:
    """
    Extract resume bullets from raw text.
    
    Uses pattern matching to identify bullet points, with focus on
    experience sections and filtering out non-work-related content.
    
    Args:
        text: Raw text from resume
        
    Returns:
        List[str]: Extracted bullets
    """
    lines = text.split('\n')
    bullets = []
    in_experience_section = False
    section_depth = 0
    
    for i, line in enumerate(lines):
        line = line.strip()
        
        if not line:
            continue
        
        # Check if this is an experience section header
        line_lower = line.lower()
        if any(re.search(pattern, line_lower) for pattern in EXPERIENCE_SECTION_PATTERNS):
            in_experience_section = True
            section_depth = 0
            continue
        
        # Check if we're leaving experience section (new major section)
        if line.isupper() and len(line) > 3 and not any(c.isdigit() for c in line):
            # Likely a section header
            if in_experience_section and section_depth > 5:
                in_experience_section = False
            section_depth = 0
            continue
        
        # Check if line matches bullet pattern
        is_bullet = any(re.match(pattern, line) for pattern in BULLET_PATTERNS)
        
        # Also consider indented lines in experience section as potential bullets
        if in_experience_section and len(line) > 20 and line[0].isupper():
            is_bullet = True
        
        if is_bullet:
            section_depth += 1
            
            # Check exclusion patterns
            line_lower = line.lower()
            if any(re.search(pattern, line_lower) for pattern in EXCLUDE_PATTERNS):
                continue
            
            # Clean up bullet
            cleaned = clean_bullet(line)
            
            if cleaned and len(cleaned) > 15:  # Minimum reasonable bullet length
                bullets.append(cleaned)
    
    # If we didn't find bullets in experience section, try to find any bullets
    if not bullets:
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            is_bullet = any(re.match(pattern, line) for pattern in BULLET_PATTERNS)
            
            if is_bullet:
                line_lower = line.lower()
                if any(re.search(pattern, line_lower) for pattern in EXCLUDE_PATTERNS):
                    continue
                
                cleaned = clean_bullet(line)
                if cleaned and len(cleaned) > 15:
                    bullets.append(cleaned)
    
    # Deduplicate while preserving order
    seen = set()
    unique_bullets = []
    for bullet in bullets:
        if bullet not in seen:
            seen.add(bullet)
            unique_bullets.append(bullet)
    
    return unique_bullets


def clean_bullet(bullet: str) -> str:
    """
    Clean and normalize a bullet point.
    
    Args:
        bullet: Raw bullet text
        
    Returns:
        str: Cleaned bullet
    """
    # Remove bullet markers
    for pattern in BULLET_PATTERNS:
        bullet = re.sub(pattern, '', bullet)
    
    # Remove extra whitespace
    bullet = ' '.join(bullet.split())
    
    # Ensure it starts with capital letter
    bullet = bullet.strip()
    if bullet and bullet[0].islower():
        bullet = bullet[0].upper() + bullet[1:]
    
    return bullet

