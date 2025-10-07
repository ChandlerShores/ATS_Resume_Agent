"""Convert DOCX resume files to PDF and TXT formats for testing."""

import sys
from pathlib import Path

# Add parent directory to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document


def convert_docx_to_txt(docx_path: Path, txt_path: Path):
    """
    Convert DOCX to plain TXT format.
    
    Args:
        docx_path: Path to DOCX file
        txt_path: Path for output TXT file
    """
    doc = Document(docx_path)
    
    # Extract all paragraphs
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)
    
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)
    
    # Write to TXT file
    with open(txt_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(paragraphs))
    
    print(f"  ✅ TXT: {txt_path.name}")


def convert_docx_to_pdf_windows(docx_path: Path, pdf_path: Path):
    """
    Convert DOCX to PDF on Windows using docx2pdf.
    
    Args:
        docx_path: Path to DOCX file
        pdf_path: Path for output PDF file
    """
    try:
        from docx2pdf import convert
        convert(str(docx_path), str(pdf_path))
        print(f"  ✅ PDF: {pdf_path.name}")
        return True
    except ImportError:
        print(f"  ⚠️  PDF: docx2pdf not installed (pip install docx2pdf)")
        return False
    except Exception as e:
        print(f"  ❌ PDF: Failed - {str(e)}")
        return False


def main():
    """Main conversion entry point."""
    print("\n" + "="*70)
    print("DOCX RESUME CONVERTER")
    print("="*70)
    print("Converting DOCX files to PDF and TXT formats for testing\n")
    
    # Source directory
    source_dir = Path("tests/resumes")
    if not source_dir.exists():
        print(f"❌ Error: Source directory not found: {source_dir}")
        sys.exit(1)
    
    # Get all DOCX files
    docx_files = list(source_dir.glob("*.docx"))
    if not docx_files:
        print(f"❌ Error: No DOCX files found in {source_dir}")
        sys.exit(1)
    
    print(f"Found {len(docx_files)} DOCX files to convert\n")
    
    # Create output directories
    pdf_dir = source_dir / "pdf"
    txt_dir = source_dir / "txt"
    pdf_dir.mkdir(exist_ok=True)
    txt_dir.mkdir(exist_ok=True)
    
    # Convert each file
    pdf_success = 0
    txt_success = 0
    
    for docx_file in docx_files:
        print(f"Converting: {docx_file.name}")
        
        # Convert to TXT
        txt_path = txt_dir / docx_file.with_suffix('.txt').name
        try:
            convert_docx_to_txt(docx_file, txt_path)
            txt_success += 1
        except Exception as e:
            print(f"  ❌ TXT: Failed - {str(e)}")
        
        # Convert to PDF (Windows-specific)
        pdf_path = pdf_dir / docx_file.with_suffix('.pdf').name
        if convert_docx_to_pdf_windows(docx_file, pdf_path):
            pdf_success += 1
        
        print()
    
    # Summary
    print("="*70)
    print("CONVERSION COMPLETE")
    print("="*70)
    print(f"TXT files: {txt_success}/{len(docx_files)} created in {txt_dir}")
    print(f"PDF files: {pdf_success}/{len(docx_files)} created in {pdf_dir}")
    print()
    
    if pdf_success == 0:
        print("⚠️  PDF conversion failed. To enable PDF conversion:")
        print("   pip install docx2pdf")
        print()
    
    print("You can now test parsing with:")
    print(f"  python scripts/parse_resume.py --resume {txt_dir}/[file].txt --role \"...\" --preview")
    print(f"  python scripts/parse_resume.py --resume {pdf_dir}/[file].pdf --role \"...\" --preview")
    print("="*70 + "\n")


if __name__ == "__main__":
    main()

